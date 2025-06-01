#Step1
#pip install pyperclip python-docs pillow pyinstaller
#Step2
#pyinstaller --onefile --windowed NotesC.py

import time
import tkinter as tk
from tkinter import filedialog, messagebox
import threading
import os
import pyperclip
from PIL import ImageGrab
from docx import Document
from docx.shared import Inches
import io
import re
import logging
from datetime import datetime

class ClipboardToWordApp:
    def __init__(self):
        self.doc_path = None
        self.last_text = ""
        self.last_image_data = None
        self.running = True
        self.last_autosave = time.time()

        # Logging setup
        logging.basicConfig(
            filename="clipboard_notes.log",
            level=logging.INFO,
            format="%(asctime)s - %(levelname)s - %(message)s"
        )

        self.start_gui()

    def start_gui(self):
        def select_file():
            file_path = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word Document", "*.docx")]
            )
            if file_path:
                self.doc_path = file_path
                if not os.path.exists(file_path):
                    try:
                        doc = Document()
                        doc.save(file_path)
                        logging.info(f"Created new document: {file_path}")
                    except Exception as e:
                        messagebox.showerror("Error", f"Could not create file:\n{e}")
                        return
                messagebox.showinfo("File Set", f"Word document set to:\n{self.doc_path}")

        def on_close():
            self.running = False
            root.destroy()
            logging.info("Application closed by user.")

        root = tk.Tk()
        root.title("Clipboard to Word Notes")
        root.geometry("300x140")
        root.resizable(False, False)

        tk.Label(root, text="Clipboard Notes to Word").pack(pady=10)
        tk.Button(root, text="Set Word File", command=select_file).pack(pady=5)
        tk.Button(root, text="Exit", command=on_close).pack(pady=5)

        threading.Thread(target=self.monitor_clipboard, daemon=True).start()
        root.mainloop()

    def monitor_clipboard(self):
        logging.info("Started clipboard monitoring.")
        while self.running:
            try:
                # Check for text
                text = pyperclip.paste()
                if text != self.last_text and isinstance(text, str) and len(text.strip()) > 0:
                    self.last_text = text
                    if self.doc_path:
                        self.append_text_to_doc(text)

                # Check for image
                image = ImageGrab.grabclipboard()
                if image and image != self.last_image_data:
                    self.last_image_data = image
                    if self.doc_path:
                        self.append_image_to_doc(image)

            except Exception as e:
                logging.error(f"Clipboard monitoring error: {e}")

            time.sleep(0.5)

    def append_text_to_doc(self, text):
        try:
            text = text.strip()
            text = re.sub(r'\n{3,}', '\n\n', text)
            if text:
                doc = Document(self.doc_path)
                doc.add_paragraph(text)
                doc.add_paragraph("")  # One blank paragraph
                doc.save(self.doc_path)
                logging.info("Text appended to document.")
        except Exception as e:
            logging.error(f"Failed to append text: {e}")
        self.auto_backup()

    def append_image_to_doc(self, image):
        try:
            stream = io.BytesIO()
            image.save(stream, format='PNG')
            stream.seek(0)

            doc = Document(self.doc_path)
            doc.add_picture(stream, width=Inches(5.5))
            doc.add_paragraph("")  # One blank paragraph
            doc.save(self.doc_path)
            logging.info("Image appended to document.")
        except Exception as e:
            logging.error(f"Failed to append image: {e}")
        self.auto_backup()

    def auto_backup(self):
        try:
            if time.time() - self.last_autosave > 600:  # 10 minutes
                backup_path = self.doc_path.replace(".docx", "_backup.docx")
                doc = Document(self.doc_path)
                doc.save(backup_path)
                self.last_autosave = time.time()
                logging.info(f"Backup saved: {backup_path}")
        except Exception as e:
            logging.error(f"Failed to create backup: {e}")

if __name__ == "__main__":
    ClipboardToWordApp()
